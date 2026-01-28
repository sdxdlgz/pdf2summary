"""
Unit tests for the Document Processor.

Tests cover:
- Bilingual Markdown creation with paragraph alignment
- Markdown to DOCX conversion
- Handling of various Markdown formatting elements
"""

import pytest
from io import BytesIO

from docx import Document

from backend.services.document_processor import (
    DocumentProcessor,
    DocumentProcessorError,
    split_into_paragraphs,
    merge_paragraphs,
)


class TestSplitIntoParagraphs:
    """Tests for the split_into_paragraphs function."""
    
    def test_empty_content(self):
        """Test splitting empty content."""
        assert split_into_paragraphs("") == []
        assert split_into_paragraphs(None) == []
    
    def test_single_paragraph(self):
        """Test splitting content with a single paragraph."""
        content = "This is a single paragraph."
        result = split_into_paragraphs(content)
        assert result == ["This is a single paragraph."]
    
    def test_multiple_paragraphs(self):
        """Test splitting content with multiple paragraphs."""
        content = "First paragraph.\n\nSecond paragraph.\n\nThird paragraph."
        result = split_into_paragraphs(content)
        assert result == ["First paragraph.", "Second paragraph.", "Third paragraph."]
    
    def test_whitespace_only_paragraphs(self):
        """Test that whitespace-only paragraphs are filtered out."""
        content = "First paragraph.\n\n   \n\nSecond paragraph."
        result = split_into_paragraphs(content)
        assert result == ["First paragraph.", "Second paragraph."]
    
    def test_paragraphs_with_extra_newlines(self):
        """Test handling of extra newlines."""
        content = "First paragraph.\n\n\n\nSecond paragraph."
        result = split_into_paragraphs(content)
        # Extra newlines create empty paragraphs which are filtered
        assert len(result) == 2
        assert "First paragraph." in result
        assert "Second paragraph." in result
    
    def test_paragraphs_with_leading_trailing_whitespace(self):
        """Test that leading/trailing whitespace is stripped."""
        content = "  First paragraph.  \n\n  Second paragraph.  "
        result = split_into_paragraphs(content)
        assert result == ["First paragraph.", "Second paragraph."]


class TestMergeParagraphs:
    """Tests for the merge_paragraphs function."""
    
    def test_empty_list(self):
        """Test merging empty list."""
        assert merge_paragraphs([]) == ""
    
    def test_single_paragraph(self):
        """Test merging single paragraph."""
        result = merge_paragraphs(["Single paragraph."])
        assert result == "Single paragraph."
    
    def test_multiple_paragraphs(self):
        """Test merging multiple paragraphs."""
        paragraphs = ["First.", "Second.", "Third."]
        result = merge_paragraphs(paragraphs)
        assert result == "First.\n\nSecond.\n\nThird."


class TestDocumentProcessorBilingualMarkdown:
    """Tests for the create_bilingual_markdown method."""
    
    @pytest.fixture
    def processor(self):
        """Create a DocumentProcessor instance."""
        return DocumentProcessor()
    
    @pytest.mark.asyncio
    async def test_empty_content(self, processor):
        """Test with empty content."""
        result = await processor.create_bilingual_markdown("", "")
        assert result == ""
    
    @pytest.mark.asyncio
    async def test_only_original(self, processor):
        """Test with only original content."""
        result = await processor.create_bilingual_markdown("Original text.", "")
        assert result == "Original text."
    
    @pytest.mark.asyncio
    async def test_only_translated(self, processor):
        """Test with only translated content."""
        result = await processor.create_bilingual_markdown("", "Translated text.")
        assert result == "Translated text."
    
    @pytest.mark.asyncio
    async def test_single_paragraph_alignment(self, processor):
        """Test alignment of single paragraphs."""
        original = "This is the original text."
        translated = "这是翻译后的文本。"
        
        result = await processor.create_bilingual_markdown(original, translated)
        
        # Should contain both original and translated
        assert "This is the original text." in result
        assert "> 这是翻译后的文本。" in result
    
    @pytest.mark.asyncio
    async def test_multiple_paragraph_alignment(self, processor):
        """Test alignment of multiple paragraphs."""
        original = "First paragraph.\n\nSecond paragraph."
        translated = "第一段。\n\n第二段。"
        
        result = await processor.create_bilingual_markdown(original, translated)
        
        # Should contain all paragraphs
        assert "First paragraph." in result
        assert "> 第一段。" in result
        assert "Second paragraph." in result
        assert "> 第二段。" in result
        
        # Should have separators between paragraph pairs
        assert "---" in result
    
    @pytest.mark.asyncio
    async def test_mismatched_paragraph_count(self, processor):
        """Test handling of mismatched paragraph counts."""
        original = "First.\n\nSecond.\n\nThird."
        translated = "第一。\n\n第二。"
        
        result = await processor.create_bilingual_markdown(original, translated)
        
        # Should handle gracefully - include all original paragraphs
        assert "First." in result
        assert "Second." in result
        assert "Third." in result
        assert "> 第一。" in result
        assert "> 第二。" in result
    
    @pytest.mark.asyncio
    async def test_blockquote_formatting(self, processor):
        """Test that translations are formatted as blockquotes."""
        original = "Original text."
        translated = "翻译文本。"
        
        result = await processor.create_bilingual_markdown(original, translated)
        
        # Translation should be prefixed with >
        assert "> 翻译文本。" in result
    
    @pytest.mark.asyncio
    async def test_multiline_paragraph_blockquote(self, processor):
        """Test blockquote formatting for multi-line paragraphs."""
        original = "Original text."
        translated = "Line one.\nLine two."
        
        result = await processor.create_bilingual_markdown(original, translated)
        
        # Each line of translation should be prefixed with >
        assert "> Line one." in result
        assert "> Line two." in result


class TestDocumentProcessorMarkdownToDocx:
    """Tests for the markdown_to_docx method."""
    
    @pytest.fixture
    def processor(self):
        """Create a DocumentProcessor instance."""
        return DocumentProcessor()
    
    def _read_docx(self, docx_bytes: bytes) -> Document:
        """Helper to read DOCX bytes into a Document object."""
        return Document(BytesIO(docx_bytes))
    
    @pytest.mark.asyncio
    async def test_empty_content(self, processor):
        """Test conversion of empty content."""
        result = await processor.markdown_to_docx("")
        
        # Should return valid DOCX bytes
        assert isinstance(result, bytes)
        assert len(result) > 0
        
        # Should be readable as DOCX
        doc = self._read_docx(result)
        assert len(doc.paragraphs) == 0
    
    @pytest.mark.asyncio
    async def test_plain_text(self, processor):
        """Test conversion of plain text."""
        content = "This is plain text."
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        assert len(doc.paragraphs) >= 1
        
        # Find the paragraph with our text
        text_found = any("This is plain text." in p.text for p in doc.paragraphs)
        assert text_found
    
    @pytest.mark.asyncio
    async def test_header_conversion(self, processor):
        """Test conversion of Markdown headers."""
        content = "# Heading 1\n\n## Heading 2\n\n### Heading 3"
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that headings are present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Heading 1" in all_text
        assert "Heading 2" in all_text
        assert "Heading 3" in all_text
    
    @pytest.mark.asyncio
    async def test_bold_text(self, processor):
        """Test conversion of bold text."""
        content = "This is **bold** text."
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that text is present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "bold" in all_text
        
        # Check for bold formatting
        bold_found = False
        for para in doc.paragraphs:
            for run in para.runs:
                if "bold" in run.text and run.bold:
                    bold_found = True
                    break
        assert bold_found
    
    @pytest.mark.asyncio
    async def test_italic_text(self, processor):
        """Test conversion of italic text."""
        content = "This is *italic* text."
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that text is present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "italic" in all_text
        
        # Check for italic formatting
        italic_found = False
        for para in doc.paragraphs:
            for run in para.runs:
                if "italic" in run.text and run.italic:
                    italic_found = True
                    break
        assert italic_found
    
    @pytest.mark.asyncio
    async def test_unordered_list(self, processor):
        """Test conversion of unordered lists."""
        content = "- Item 1\n- Item 2\n- Item 3"
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that list items are present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Item 1" in all_text
        assert "Item 2" in all_text
        assert "Item 3" in all_text
        
        # Check for bullet character
        assert "•" in all_text
    
    @pytest.mark.asyncio
    async def test_ordered_list(self, processor):
        """Test conversion of ordered lists."""
        content = "1. First\n2. Second\n3. Third"
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that list items are present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "First" in all_text
        assert "Second" in all_text
        assert "Third" in all_text
    
    @pytest.mark.asyncio
    async def test_blockquote(self, processor):
        """Test conversion of blockquotes."""
        content = "> This is a quote."
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that quote text is present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "This is a quote." in all_text
    
    @pytest.mark.asyncio
    async def test_horizontal_rule(self, processor):
        """Test conversion of horizontal rules."""
        content = "Before\n\n---\n\nAfter"
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that content around the rule is present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Before" in all_text
        assert "After" in all_text
    
    @pytest.mark.asyncio
    async def test_inline_code(self, processor):
        """Test conversion of inline code."""
        content = "Use the `print()` function."
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that code text is present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "print()" in all_text
    
    @pytest.mark.asyncio
    async def test_link_conversion(self, processor):
        """Test conversion of links."""
        content = "Visit [Google](https://google.com) for search."
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that link text and URL are present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Google" in all_text
        assert "https://google.com" in all_text
    
    @pytest.mark.asyncio
    async def test_image_placeholder(self, processor):
        """Test that images are converted to placeholders."""
        content = "![Alt text](image.png)"
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that image placeholder is present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Image:" in all_text or "Alt text" in all_text
    
    @pytest.mark.asyncio
    async def test_complex_document(self, processor):
        """Test conversion of a complex document with multiple elements."""
        content = """# Document Title

This is an introduction paragraph with **bold** and *italic* text.

## Section 1

Here is a list:

- Item one
- Item two
- Item three

## Section 2

> This is a blockquote with important information.

And here is some `inline code` in a paragraph.

---

Final paragraph."""
        
        result = await processor.markdown_to_docx(content)
        
        doc = self._read_docx(result)
        
        # Check that all major elements are present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Document Title" in all_text
        assert "introduction paragraph" in all_text
        assert "Section 1" in all_text
        assert "Item one" in all_text
        assert "Section 2" in all_text
        assert "blockquote" in all_text
        assert "inline code" in all_text
        assert "Final paragraph" in all_text
    
    @pytest.mark.asyncio
    async def test_bilingual_document_to_docx(self, processor):
        """Test converting a bilingual Markdown document to DOCX."""
        # First create a bilingual document
        original = "Hello, world!\n\nThis is a test."
        translated = "你好，世界！\n\n这是一个测试。"
        
        bilingual_md = await processor.create_bilingual_markdown(original, translated)
        
        # Then convert to DOCX
        result = await processor.markdown_to_docx(bilingual_md)
        
        doc = self._read_docx(result)
        
        # Check that both original and translated content is present
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Hello, world!" in all_text
        assert "你好，世界！" in all_text
        assert "This is a test." in all_text
        assert "这是一个测试。" in all_text
