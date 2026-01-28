"""
Document Processor for the Research Report Processor.

This module provides document processing functionality including:
- Creating bilingual Markdown documents with paragraph-by-paragraph alignment
- Converting Markdown to DOCX format

Requirements:
- 5.5: Produce bilingual document with original and translated text aligned paragraph by paragraph
- 5.6: Create both MD and DOCX format files
- 7.3: Provide bilingual translation MD file for download
- 7.4: Provide bilingual translation DOCX file for download
"""

import io
import logging
import re
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


logger = logging.getLogger(__name__)


# Paragraph separator for splitting content
PARAGRAPH_SEPARATOR = "\n\n"

# Markdown patterns for parsing
HEADER_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$')
BOLD_PATTERN = re.compile(r'\*\*(.+?)\*\*')
ITALIC_PATTERN = re.compile(r'\*(.+?)\*')
BOLD_ITALIC_PATTERN = re.compile(r'\*\*\*(.+?)\*\*\*')
UNORDERED_LIST_PATTERN = re.compile(r'^[-*+]\s+(.+)$')
ORDERED_LIST_PATTERN = re.compile(r'^(\d+)\.\s+(.+)$')
CODE_INLINE_PATTERN = re.compile(r'`([^`]+)`')
LINK_PATTERN = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
IMAGE_PATTERN = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')


class DocumentProcessorError(Exception):
    """
    Exception raised when document processing operations fail.
    
    Attributes:
        message: Human-readable error description
        original_error: The underlying error that caused the failure
    """
    
    def __init__(
        self,
        message: str,
        original_error: Optional[Exception] = None,
    ):
        self.message = message
        self.original_error = original_error
        super().__init__(self.message)


def split_into_paragraphs(content: str) -> list[str]:
    """
    Split content into paragraphs by double newlines.
    
    Args:
        content: The content to split
        
    Returns:
        List of non-empty paragraphs
    """
    if not content:
        return []
    
    # Split by double newlines (paragraph separator)
    paragraphs = content.split(PARAGRAPH_SEPARATOR)
    
    # Filter out empty paragraphs and strip whitespace
    paragraphs = [p.strip() for p in paragraphs if p.strip()]
    
    return paragraphs


def merge_paragraphs(paragraphs: list[str]) -> str:
    """
    Merge paragraphs back into a single document.
    
    Args:
        paragraphs: List of paragraphs
        
    Returns:
        Merged content string
    """
    return PARAGRAPH_SEPARATOR.join(paragraphs)


class DocumentProcessor:
    """
    Document processor for creating bilingual documents and format conversion.
    
    This class provides:
    - Bilingual Markdown creation with paragraph-by-paragraph alignment
    - Markdown to DOCX conversion with basic formatting support
    
    The bilingual document format interleaves original and translated paragraphs,
    making it easy to compare the original text with its translation.
    
    Validates: Requirements 5.5, 5.6, 7.3, 7.4
    """
    
    def __init__(self):
        """Initialize the document processor."""
        pass
    
    async def create_bilingual_markdown(
        self,
        original: str,
        translated: str,
    ) -> str:
        """
        Create bilingual Markdown with original and translated aligned paragraph by paragraph.
        
        This method:
        1. Splits both original and translated content by paragraphs (double newline)
        2. Interleaves original and translated paragraphs
        3. Adds visual separators between paragraph pairs for readability
        
        The output format is:
        ```
        [Original Paragraph 1]
        
        > [Translated Paragraph 1]
        
        ---
        
        [Original Paragraph 2]
        
        > [Translated Paragraph 2]
        
        ---
        ...
        ```
        
        Args:
            original: Original Markdown content
            translated: Translated Markdown content
            
        Returns:
            Bilingual Markdown with aligned paragraphs
            
        Validates: Requirement 5.5 - Bilingual document with paragraph alignment
        """
        if not original and not translated:
            return ""
        
        if not original:
            return translated
        
        if not translated:
            return original
        
        # Split into paragraphs
        original_paragraphs = split_into_paragraphs(original)
        translated_paragraphs = split_into_paragraphs(translated)
        
        logger.debug(
            "Creating bilingual document: %d original paragraphs, %d translated paragraphs",
            len(original_paragraphs),
            len(translated_paragraphs),
        )
        
        # Handle mismatched paragraph counts
        if len(original_paragraphs) != len(translated_paragraphs):
            logger.warning(
                "Paragraph count mismatch: original=%d, translated=%d",
                len(original_paragraphs),
                len(translated_paragraphs),
            )
        
        # Interleave paragraphs
        bilingual_parts: list[str] = []
        max_paragraphs = max(len(original_paragraphs), len(translated_paragraphs))
        
        for i in range(max_paragraphs):
            # Get original paragraph (or empty if exhausted)
            orig_para = original_paragraphs[i] if i < len(original_paragraphs) else ""
            
            # Get translated paragraph (or empty if exhausted)
            trans_para = translated_paragraphs[i] if i < len(translated_paragraphs) else ""
            
            # Add original paragraph
            if orig_para:
                bilingual_parts.append(orig_para)
            
            # Add translated paragraph with blockquote formatting for visual distinction
            if trans_para:
                # Format translation as blockquote for visual distinction
                # Handle multi-line paragraphs by prefixing each line with >
                trans_lines = trans_para.split('\n')
                quoted_trans = '\n'.join(f"> {line}" for line in trans_lines)
                bilingual_parts.append(quoted_trans)
            
            # Add separator between paragraph pairs (except for the last pair)
            if i < max_paragraphs - 1:
                bilingual_parts.append("---")
        
        result = PARAGRAPH_SEPARATOR.join(bilingual_parts)
        
        logger.info(
            "Created bilingual document: %d paragraph pairs",
            max_paragraphs,
        )
        
        return result
    
    async def markdown_to_docx(self, markdown_content: str) -> bytes:
        """
        Convert Markdown to DOCX format.
        
        This method handles basic Markdown formatting:
        - Headers (# to ######)
        - Bold (**text**)
        - Italic (*text*)
        - Unordered lists (-, *, +)
        - Ordered lists (1., 2., etc.)
        - Blockquotes (> text)
        - Horizontal rules (---)
        - Inline code (`code`)
        - Links [text](url)
        
        Args:
            markdown_content: Markdown content to convert
            
        Returns:
            DOCX file content as bytes
            
        Raises:
            DocumentProcessorError: If conversion fails
            
        Validates: Requirements 5.6, 7.4 - Create DOCX format files
        """
        if not markdown_content:
            # Return empty DOCX
            doc = Document()
            return self._document_to_bytes(doc)
        
        try:
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            # Split content into paragraphs
            paragraphs = split_into_paragraphs(markdown_content)
            
            for para_text in paragraphs:
                self._add_markdown_paragraph(doc, para_text)
            
            logger.info(
                "Converted Markdown to DOCX: %d paragraphs",
                len(paragraphs),
            )
            
            return self._document_to_bytes(doc)
        
        except Exception as e:
            logger.error("Failed to convert Markdown to DOCX: %s", e)
            raise DocumentProcessorError(
                message=f"Failed to convert Markdown to DOCX: {e}",
                original_error=e,
            ) from e
    
    def _document_to_bytes(self, doc: Document) -> bytes:
        """
        Convert a Document object to bytes.
        
        Args:
            doc: The python-docx Document object
            
        Returns:
            DOCX file content as bytes
        """
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _add_markdown_paragraph(self, doc: Document, para_text: str) -> None:
        """
        Add a Markdown paragraph to the DOCX document.
        
        Handles various Markdown elements:
        - Headers
        - Lists
        - Blockquotes
        - Horizontal rules
        - Regular paragraphs with inline formatting
        
        Args:
            doc: The python-docx Document object
            para_text: The paragraph text in Markdown format
        """
        # Handle multi-line paragraphs (e.g., multi-line blockquotes)
        lines = para_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for horizontal rule
            if line in ('---', '***', '___'):
                # Add a horizontal line (using a paragraph with bottom border)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(12)
                # Add a thin line using underscores
                run = p.add_run('_' * 50)
                run.font.color.rgb = None  # Use default color
                continue
            
            # Check for header
            header_match = HEADER_PATTERN.match(line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                self._add_header(doc, text, level)
                continue
            
            # Check for blockquote
            if line.startswith('>'):
                # Remove the > prefix and any leading space
                quote_text = line[1:].lstrip()
                self._add_blockquote(doc, quote_text)
                continue
            
            # Check for unordered list
            ul_match = UNORDERED_LIST_PATTERN.match(line)
            if ul_match:
                list_text = ul_match.group(1)
                self._add_list_item(doc, list_text, ordered=False)
                continue
            
            # Check for ordered list
            ol_match = ORDERED_LIST_PATTERN.match(line)
            if ol_match:
                list_text = ol_match.group(2)
                self._add_list_item(doc, list_text, ordered=True, number=int(ol_match.group(1)))
                continue
            
            # Regular paragraph with inline formatting
            self._add_formatted_paragraph(doc, line)
    
    def _add_header(self, doc: Document, text: str, level: int) -> None:
        """
        Add a header to the document.
        
        Args:
            doc: The python-docx Document object
            text: The header text
            level: Header level (1-6)
        """
        # Map Markdown header levels to Word heading styles
        # Word has Heading 1-9, we use 1-6
        heading_level = min(level, 6)
        
        # Add heading
        heading = doc.add_heading(level=heading_level)
        self._add_formatted_text(heading, text)
    
    def _add_blockquote(self, doc: Document, text: str) -> None:
        """
        Add a blockquote to the document.
        
        Args:
            doc: The python-docx Document object
            text: The blockquote text
        """
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        # Add italic formatting for blockquotes
        self._add_formatted_text(p, text, default_italic=True)
    
    def _add_list_item(
        self,
        doc: Document,
        text: str,
        ordered: bool = False,
        number: int = 1,
    ) -> None:
        """
        Add a list item to the document.
        
        Args:
            doc: The python-docx Document object
            text: The list item text
            ordered: Whether this is an ordered list
            number: The item number (for ordered lists)
        """
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        
        # Add bullet or number
        if ordered:
            prefix = f"{number}. "
        else:
            prefix = "• "
        
        p.add_run(prefix)
        self._add_formatted_text(p, text)
    
    def _add_formatted_paragraph(self, doc: Document, text: str) -> None:
        """
        Add a regular paragraph with inline formatting.
        
        Args:
            doc: The python-docx Document object
            text: The paragraph text with Markdown formatting
        """
        p = doc.add_paragraph()
        self._add_formatted_text(p, text)
    
    def _add_formatted_text(
        self,
        paragraph,
        text: str,
        default_italic: bool = False,
    ) -> None:
        """
        Add text with inline Markdown formatting to a paragraph.
        
        Handles:
        - Bold (**text**)
        - Italic (*text*)
        - Bold+Italic (***text***)
        - Inline code (`code`)
        - Links [text](url)
        
        Args:
            paragraph: The python-docx Paragraph object
            text: The text with Markdown formatting
            default_italic: Whether to apply italic by default (for blockquotes)
        """
        # Process the text to handle inline formatting
        # We'll use a simple state machine approach
        
        # First, handle images by removing them (DOCX image handling is complex)
        text = IMAGE_PATTERN.sub(r'[Image: \1]', text)
        
        # Handle links by converting to text with URL
        text = LINK_PATTERN.sub(r'\1 (\2)', text)
        
        # Now process bold, italic, and code
        # We need to handle nested formatting carefully
        
        # Split text into segments with formatting markers
        segments = self._parse_inline_formatting(text)
        
        for segment_text, is_bold, is_italic, is_code in segments:
            if not segment_text:
                continue
            
            run = paragraph.add_run(segment_text)
            
            # Apply formatting
            if is_bold:
                run.bold = True
            if is_italic or default_italic:
                run.italic = True
            if is_code:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
    
    def _parse_inline_formatting(
        self,
        text: str,
    ) -> list[tuple[str, bool, bool, bool]]:
        """
        Parse text for inline Markdown formatting.
        
        Returns a list of tuples: (text, is_bold, is_italic, is_code)
        
        Args:
            text: The text to parse
            
        Returns:
            List of (text, is_bold, is_italic, is_code) tuples
        """
        segments: list[tuple[str, bool, bool, bool]] = []
        
        # Pattern to match all inline formatting
        # Order matters: check bold+italic first, then bold, then italic, then code
        combined_pattern = re.compile(
            r'(\*\*\*(.+?)\*\*\*)|'  # Bold+Italic
            r'(\*\*(.+?)\*\*)|'      # Bold
            r'(\*(.+?)\*)|'          # Italic
            r'(`([^`]+)`)'           # Code
        )
        
        last_end = 0
        
        for match in combined_pattern.finditer(text):
            # Add any text before this match as plain text
            if match.start() > last_end:
                plain_text = text[last_end:match.start()]
                if plain_text:
                    segments.append((plain_text, False, False, False))
            
            # Determine which group matched
            if match.group(2):  # Bold+Italic
                segments.append((match.group(2), True, True, False))
            elif match.group(4):  # Bold
                segments.append((match.group(4), True, False, False))
            elif match.group(6):  # Italic
                segments.append((match.group(6), False, True, False))
            elif match.group(8):  # Code
                segments.append((match.group(8), False, False, True))
            
            last_end = match.end()
        
        # Add any remaining text after the last match
        if last_end < len(text):
            remaining = text[last_end:]
            if remaining:
                segments.append((remaining, False, False, False))
        
        # If no formatting was found, return the whole text as plain
        if not segments and text:
            segments.append((text, False, False, False))
        
        return segments
